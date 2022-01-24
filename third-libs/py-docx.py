from docx import Document

document = Document()
document.add_heading('The REAL meaning of the universe',0)
document.add_heading('The role of dolphins', level=0)
document.add_heading('The role of dolphins', level=1)
document.add_heading('The role of dolphins', level=2)
document.add_heading('The role of dolphins', level=3)
document.add_heading('The role of dolphins', level=4)

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
document.save('tests/00.docx')